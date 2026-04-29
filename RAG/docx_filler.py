"""
docx_filler.py
==============
Module parse JSON fields từ PromptAPI vào template DOCX có sẵn placeholder.

Placeholder format trong template: {{FIELD_NAME}}
  VD: {{TEN_CQ_BAN_HANH}}, {{NOI_DUNG_DIEU_1}}, ...

Cấu trúc thư mục mặc định (có thể override):
  Forms/docx/          ← chứa các template .docx
  drafts/              ← output file đã điền

Cách dùng độc lập:
    from docx_filler import DocxFiller
    filler = DocxFiller()
    output_path = filler.fill(form_id="Form_02", fields={...})

Tích hợp với PromptAPI:
    result = api.draft(query, extras=extras)
    output_path = filler.fill_from_result(result)
"""

from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════
# MAPPING: form_id → tên file template
# Thêm form mới vào đây khi có template mới
# ═══════════════════════════════════════════════════════════════
FORM_TEMPLATE_MAP: Dict[str, str] = {
    "Form_01": "Mau_1.1",       # Nghị quyết cá biệt
    "Form_02": "Mau_1.2",       # Quyết định cá biệt
    "Form_03": "Mau_1.3",       # Quyết định quy phạm
    "Form_04": "Mau_1.4",       # Văn bản có tên loại
    "Form_05": "Mau_1.5",        # Công văn
    "Form_06": "Mau_1.6",       # Công điện
    "Form_07": "Mau_1.7",       # Giấy mời
    "Form_08": "Mau_1.8",     # Giấy giới thiệu
    "Form_09": "Mau_1.9",       # Biên bản
    "Form_10": "Mau_1.10",      # Giấy nghỉ phép
}


class DocxFiller:
    """
    Fill placeholder {{FIELD}} trong template DOCX bằng python-docx.
    Giữ nguyên toàn bộ formatting gốc của template.

    Parameters
    ----------
    forms_dir : str | Path
        Thư mục chứa các file template .docx. Mặc định: "Forms/docx"
    drafts_dir : str | Path
        Thư mục lưu file đã điền. Mặc định: "drafts"
    """

    def __init__(
        self,
        forms_dir : str | Path = "Forms/docx",
        drafts_dir: str | Path = "drafts",
    ):
        self.forms_dir  = Path(forms_dir)
        self.drafts_dir = Path(drafts_dir)
        self.drafts_dir.mkdir(parents=True, exist_ok=True)
        self._ensure_docx()

    # ───────────────────────────────────────────────────────────
    # PUBLIC API
    # ───────────────────────────────────────────────────────────
    def fill(
        self,
        form_id : str,
        fields  : Dict[str, str],
        out_name: Optional[str] = None,
    ) -> Path:
        """
        Điền fields vào template tương ứng với form_id.

        Parameters
        ----------
        form_id : str
            Mã form VD "Form_02". Phải có trong FORM_TEMPLATE_MAP.
        fields : dict
            Dict {FIELD_NAME: value} từ PromptAPI result["fields"].
        out_name : str, optional
            Tên file output (không cần .docx). Mặc định: auto-generate theo timestamp.

        Returns
        -------
        Path
            Đường dẫn tuyệt đối đến file đã điền trong drafts/.

        Raises
        ------
        FileNotFoundError
            Nếu không tìm thấy template.
        RuntimeError
            Nếu python-docx chưa cài.
        """
        template_path = self._resolve_template(form_id)
        output_path   = self._resolve_output(form_id, out_name)

        # Copy template → draft để không làm hỏng bản gốc
        shutil.copy2(template_path, output_path)
        logger.info(f"Template: {template_path.name} → {output_path.name}")

        filled, missing, unfilled = self._fill_docx(output_path, fields)
        self._log_summary(filled, missing, unfilled, output_path)

        return output_path

    def fill_from_result(
        self,
        result  : Dict[str, Any],
        out_name: Optional[str] = None,
    ) -> Optional[Path]:
        """
        Convenience wrapper: nhận trực tiếp result dict từ PromptAPI.draft().

        Parameters
        ----------
        result : dict
            Kết quả từ api.draft(). Phải có status="ok".
        out_name : str, optional
            Tên file output tùy chỉnh.

        Returns
        -------
        Path | None
            Path đến file đã điền, hoặc None nếu result không hợp lệ.
        """
        if result.get("status") != "ok":
            logger.warning(
                f"fill_from_result: status={result.get('status')} — bỏ qua. "
                f"Lỗi: {result.get('error', 'không có LLM output')}"
            )
            return None

        form_id = result.get("meta", {}).get("form_id", "")
        fields  = result.get("fields", {})

        if not form_id:
            logger.error("fill_from_result: meta.form_id trống — không thể xác định template.")
            return None

        if not fields:
            logger.warning("fill_from_result: fields rỗng — file draft sẽ còn placeholder.")

        return self.fill(form_id, fields, out_name=out_name)

    def list_templates(self) -> List[Dict[str, str]]:
        """Trả về danh sách template có sẵn với form_id và path."""
        result = []
        for form_id, prefix in FORM_TEMPLATE_MAP.items():
            path = self._find_template_by_prefix(prefix)
            result.append({
                "form_id" : form_id,
                "prefix"  : prefix,
                "path"    : str(path) if path else "NOT FOUND",
                "exists"  : str(path is not None),
            })
        return result

    # ───────────────────────────────────────────────────────────
    # INTERNAL: RESOLVE PATHS
    # ───────────────────────────────────────────────────────────
    def _resolve_template(self, form_id: str) -> Path:
        """Tìm file template .docx theo form_id."""
        prefix = FORM_TEMPLATE_MAP.get(form_id)
        if not prefix:
            # Fallback: tìm file có chứa form_id trong tên
            fallback = list(self.forms_dir.glob(f"*{form_id}*.docx"))
            if fallback:
                return fallback[0]
            raise FileNotFoundError(
                f"form_id '{form_id}' không có trong FORM_TEMPLATE_MAP "
                f"và không tìm thấy file phù hợp trong {self.forms_dir}"
            )

        path = self._find_template_by_prefix(prefix)
        if path is None:
            raise FileNotFoundError(
                f"Không tìm thấy template với prefix '{prefix}' trong {self.forms_dir}. "
                f"Các file hiện có: {[f.name for f in self.forms_dir.glob('*.docx')]}"
            )
        return path

    def _find_template_by_prefix(self, prefix: str) -> Optional[Path]:
        """Glob tìm file .docx bắt đầu bằng prefix (case-insensitive fallback)."""
        # Exact prefix match
        matches = list(self.forms_dir.glob(f"{prefix}*.docx"))
        if matches:
            return matches[0]
        # Case-insensitive fallback
        prefix_lower = prefix.lower()
        for f in self.forms_dir.glob("*.docx"):
            if f.name.lower().startswith(prefix_lower):
                return f
        return None

    def _resolve_output(self, form_id: str, out_name: Optional[str]) -> Path:
        """Tạo tên file output trong drafts/."""
        if out_name:
            stem = out_name.removesuffix(".docx")
        else:
            ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
            stem = f"draft_{form_id}_{ts}"
        return self.drafts_dir / f"{stem}.docx"

    # ───────────────────────────────────────────────────────────
    # INTERNAL: FILL LOGIC (python-docx)
    # ───────────────────────────────────────────────────────────
    def _fill_docx(
        self,
        docx_path: Path,
        fields   : Dict[str, str],
    ) -> Tuple[List[str], List[str], List[str]]:
        """
        Thay thế {{FIELD}} trong toàn bộ document (paragraphs + tables + headers/footers).
        Giữ nguyên formatting (bold, italic, font size, color...) của từng placeholder.

        Strategy:
          Mỗi paragraph được xử lý ở cấp độ XML run (<w:r>) để đảm bảo
          formatting của placeholder được kế thừa chính xác vào giá trị thay thế.
          Tránh việc gom text vào run[0] gây mất/sai formatting.

        Returns
        -------
        filled   : list of field names đã được fill
        missing  : fields trong JSON nhưng không có placeholder trong doc
        unfilled : placeholders còn lại trong doc chưa có giá trị
        """
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        import copy

        doc = Document(docx_path)

        # Chuẩn hóa fields: strip whitespace
        normalized = {k.strip(): (v.strip() if v else "") for k, v in fields.items()}
        filled_set: set[str] = set()

        # Regex khớp placeholder có thể bị split qua nhiều runs
        PH_PATTERN = re.compile(r'\{\{([^}]+)\}\}')

        def _get_run_text(r_elem) -> str:
            """Lấy text từ XML element <w:r>."""
            t = r_elem.find(qn('w:t'))
            return t.text or "" if t is not None else ""

        def _set_run_text(r_elem, text: str) -> None:
            """Ghi text vào <w:r>, thêm xml:space='preserve' nếu cần."""
            t = r_elem.find(qn('w:t'))
            if t is None:
                t = etree.SubElement(r_elem, qn('w:t'))
            t.text = text
            if text and (text[0] == ' ' or text[-1] == ' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            elif '{http://www.w3.org/XML/1998/namespace}space' in t.attrib:
                # Xóa preserve nếu không cần
                del t.attrib['{http://www.w3.org/XML/1998/namespace}space']

        def _clone_rpr(source_r: etree._Element) -> Optional[etree._Element]:
            """Clone <w:rPr> (run properties = formatting) từ một run."""
            rpr = source_r.find(qn('w:rPr'))
            return copy.deepcopy(rpr) if rpr is not None else None

        def _make_run_with_rpr(rpr, text: str) -> etree._Element:
            """Tạo <w:r> mới với formatting clone từ rpr và text cho trước."""
            new_r = etree.Element(qn('w:r'))
            if rpr is not None:
                new_r.append(copy.deepcopy(rpr))
            t_elem = etree.SubElement(new_r, qn('w:t'))
            t_elem.text = text
            if text and (text[0] == ' ' or text[-1] == ' '):
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return new_r

        def replace_in_paragraph(para) -> None:
            """
            Xử lý một paragraph:
            1. Gom text của tất cả runs, track vị trí char → run index.
            2. Tìm placeholder trong chuỗi gom.
            3. Với mỗi placeholder:
               - Xác định run nào chứa phần lớn nhất của placeholder (→ donor run).
               - Clone rPr (formatting) từ donor run đó.
               - Tách paragraph thành segments: [text trước PH] [value với rPr PH] [text sau PH].
               - Rebuild các w:r tương ứng, giữ rPr của từng segment.
            """
            p_elem = para._p

            # Lấy tất cả w:r trực tiếp trong paragraph (bỏ qua w:ins, w:del, v.v.)
            runs = p_elem.findall(qn('w:r'))
            if not runs:
                return

            # Bước 1: Gom text + map char_index → (run_elem, local_char_index)
            full_text = ""
            char_map: List[etree._Element] = []  # char_map[i] = run chứa ký tự i
            for r in runs:
                t = _get_run_text(r)
                for _ in t:
                    char_map.append(r)
                full_text += t

            if "{{" not in full_text:
                return

            # Bước 2: Tìm tất cả placeholder
            matches = list(PH_PATTERN.finditer(full_text))
            if not matches:
                return

            # Bước 3: Xây dựng danh sách segments
            # Mỗi segment = (text, source_run_for_rpr, is_multiline)
            segments: List[Tuple[str, etree._Element, bool]] = []
            cursor = 0

            for m in matches:
                ph_key   = m.group(1).strip()
                ph_start = m.start()
                ph_end   = m.end()

                if cursor < ph_start:
                    _append_text_segments(
                        segments, full_text[cursor:ph_start],
                        char_map[cursor:ph_start], runs
                    )

                donor_run = char_map[ph_start] if ph_start < len(char_map) else runs[-1]

                if ph_key in normalized:
                    value = normalized[ph_key]
                    filled_set.add(ph_key)

                    # ── Logic đặc biệt cho NOI_DUNG_DIEN_BIEN ──────────────
                    if ph_key == "NOI_DUNG_DIEN_BIEN":
                        value = _normalize_dien_bien(value)
                    # ────────────────────────────────────────────────────────

                    segments.append((value, donor_run, "\n" in value))
                else:
                    segments.append((m.group(0), donor_run, False))

                cursor = ph_end

            if cursor < len(full_text):
                _append_text_segments(
                    segments, full_text[cursor:],
                    char_map[cursor:], runs
                )

            if not segments:
                return

            # Bước 4: Rebuild w:r elements
            # is_multiline=True → tách thành nhiều <w:p>, mỗi \n = 1 paragraph mới
            first_run_idx = list(p_elem).index(runs[0])
            for r in runs:
                p_elem.remove(r)

            parent_elem  = p_elem.getparent()
            p_insert_idx = list(parent_elem).index(p_elem)

            insert_pos  = first_run_idx
            extra_paras: List[etree._Element] = []

            for seg_text, src_run, is_multiline in segments:
                if not seg_text:
                    continue
                rpr = _clone_rpr(src_run)

                if not is_multiline:
                    new_r = _make_run_with_rpr(rpr, seg_text)
                    p_elem.insert(insert_pos, new_r)
                    insert_pos += 1
                else:
                    lines = seg_text.split("\n")
                    for line_idx, line in enumerate(lines):
                        if line_idx == 0:
                            if line:
                                new_r = _make_run_with_rpr(rpr, line)
                                p_elem.insert(insert_pos, new_r)
                                insert_pos += 1
                        else:
                            new_p = etree.Element(qn('w:p'))
                            pPr   = p_elem.find(qn('w:pPr'))
                            if pPr is not None:
                                new_p.append(copy.deepcopy(pPr))
                            if line:
                                new_r = _make_run_with_rpr(rpr, line)
                                new_p.append(new_r)
                            extra_paras.append(new_p)

            for offset, new_p in enumerate(extra_paras, start=1):
                parent_elem.insert(p_insert_idx + offset, new_p)

        def _append_text_segments(
            segments: List[Tuple[str, etree._Element]],
            text    : str,
            run_map : List[etree._Element],
            runs    : List[etree._Element],
        ) -> None:
            """
            Chia đoạn text thành sub-segments theo run nguồn để giữ formatting.
            Các ký tự liên tiếp thuộc cùng một run được gom thành 1 segment.
            """
            if not text or not run_map:
                return
            i = 0
            while i < len(text):
                src = run_map[i]
                j = i + 1
                while j < len(text) and run_map[j] is src:
                    j += 1
                segments.append((text[i:j], src, False))
                i = j

        def process_paragraphs(paragraphs) -> None:
            for para in paragraphs:
                replace_in_paragraph(para)

        def process_tables(tables) -> None:
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)
                        process_tables(cell.tables)

        # Body
        process_paragraphs(doc.paragraphs)
        process_tables(doc.tables)

        # Headers & Footers (tất cả sections)
        for section in doc.sections:
            for hf in [
                section.header, section.footer,
                section.even_page_header, section.even_page_footer,
                section.first_page_header, section.first_page_footer,
            ]:
                if hf is not None:
                    process_paragraphs(hf.paragraphs)
                    process_tables(hf.tables)

        doc.save(docx_path)

        # Tổng hợp kết quả
        missing  = [k for k in normalized if k not in filled_set and normalized[k]]
        unfilled = _scan_unfilled(docx_path)

        return sorted(filled_set), missing, unfilled

    # ───────────────────────────────────────────────────────────
    # INTERNAL: UTILS
    # ───────────────────────────────────────────────────────────
    def _ensure_docx(self) -> None:
        """Kiểm tra python-docx đã cài chưa."""
        try:
            import docx  # noqa
        except ImportError:
            raise RuntimeError(
                "python-docx chưa được cài. Chạy: pip install python-docx"
            )

    def _log_summary(
        self,
        filled  : List[str],
        missing : List[str],
        unfilled: List[str],
        output  : Path,
    ) -> None:
        logger.info(f"✅ Đã điền {len(filled)} fields → {output}")
        if missing:
            logger.warning(
                f"⚠️  {len(missing)} field(s) trong JSON nhưng không có placeholder: "
                + ", ".join(missing)
            )
        if unfilled:
            logger.warning(
                f"⚠️  {len(unfilled)} placeholder chưa có giá trị: "
                + ", ".join(unfilled)
            )


# ═══════════════════════════════════════════════════════════════
# HELPER: chuẩn hóa NOI_DUNG_DIEN_BIEN — tự nhận biết "1.", "2."...
# và chèn xuống 1 dòng giữa các mục nếu chưa có
# ═══════════════════════════════════════════════════════════════
def _normalize_dien_bien(text: str) -> str:
    """
    Nhận diện các mục đánh số "1.", "2.", ... trong text và đảm bảo
    mỗi mục bắt đầu trên một dòng mới (xuống 1 dòng, không tạo dòng trống).

    Ví dụ input (không có \\n):
        "1. Báo cáo tình hình. 2. Thảo luận. 3. Kết luận."
    Output:
        "1. Báo cáo tình hình.\\n2. Thảo luận.\\n3. Kết luận."

    Ví dụ input (đã có \\n\\n — thu gọn lại thành 1 lần):
        "1. Báo cáo\\n\\n2. Thảo luận\\n\\n3. Kết luận"
    Output:
        "1. Báo cáo\\n2. Thảo luận\\n3. Kết luận"
    """
    ITEM_PATTERN = re.compile(r'(\d+\.\s)')

    parts = ITEM_PATTERN.split(text)
    if len(parts) <= 1:
        return text

    items = []
    prefix = parts[0].strip()
    i = 1
    while i < len(parts) - 1:
        bullet  = parts[i]
        content = parts[i + 1]
        items.append(bullet + content.strip())
        i += 2

    result_parts = []
    if prefix:
        result_parts.append(prefix)
    result_parts.extend(items)

    # Nối bằng "\n" — mỗi mục xuống đúng 1 dòng
    return "\n".join(result_parts)


# ═══════════════════════════════════════════════════════════════
# HELPER: scan remaining {{...}} in saved file
# ═══════════════════════════════════════════════════════════════
def _scan_unfilled(docx_path: Path) -> List[str]:
    """Quét toàn bộ text trong docx để tìm placeholder còn sót."""
    from docx import Document
    doc      = Document(docx_path)
    found    : set[str] = set()
    pattern  = re.compile(r'\{\{([^}]+)\}\}')

    def scan_paragraphs(paras):
        for para in paras:
            text = "".join(r.text for r in para.runs)
            for m in pattern.finditer(text):
                found.add(m.group(0))

    def scan_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_paragraphs(cell.paragraphs)
                    scan_tables(cell.tables)

    scan_paragraphs(doc.paragraphs)
    scan_tables(doc.tables)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                scan_paragraphs(hf.paragraphs)

    return sorted(found)