"""
promptTemplates.py
==================
Template prompt tối ưu cho hệ thống RAG soạn thảo Văn bản Hành chính Việt Nam.

Tuân thủ Nghị định 30/2020/NĐ-CP về Công tác văn thư.

Các yêu cầu prompt được tách thành skill files trong thư mục skills/:
    skills/system_drafter.md        — System prompt soạn thảo
    skills/system_legal_qa.md       — System prompt hỏi luật
    skills/block_legal_context.md   — Hướng dẫn đọc LEGAL_CONTEXT
    skills/block_json_output.md     — Yêu cầu JSON đầu ra
    skills/block_legal_qa_answer.md — Yêu cầu trả lời hỏi luật
    skills/noi_dung_structure.md    — Cấu trúc thân văn bản theo loại

LLM output: JSON object duy nhất với 1 key duy nhất:
    - "fields": dict {FIELD_NAME: value} — toàn bộ placeholder đã điền,
                BAO GỒM cả các field NOI_DUNG_* (thân văn bản tự do).

Usage:
    from promptTemplates import build_messages, parse_llm_json, fill_word_template
    messages = build_messages(query, retrieved)
    # ... gọi LLM API, lấy response text ...
    parsed = parse_llm_json(llm_response_text)
    fill_word_template("template/Form_05.docx", parsed, "output/van_ban.docx")
"""

from __future__ import annotations

import json
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ============================================================
# 1. SKILL LOADER
# ============================================================

SKILLS_DIR = Path(__file__).parent / "skills"

_skill_cache: Dict[str, str] = {}


def _load_skill(filename: str) -> str:
    if filename in _skill_cache:
        return _skill_cache[filename]

    path = SKILLS_DIR / filename
    if not path.exists():
        raise FileNotFoundError(
            f"Skill file không tồn tại: {path}\n"
            f"Đảm bảo thư mục skills/ nằm cùng cấp với promptTemplates.py"
        )
    content = path.read_text(encoding="utf-8").strip()
    _skill_cache[filename] = content
    return content


def _skill(filename: str) -> str:
    """Shorthand wrapper cho _load_skill."""
    return _load_skill(filename)


# ============================================================
# 2. SYSTEM PROMPTS (lazy-loaded từ skill files)
# ============================================================

def get_system_prompt() -> str:
    """System prompt cho chế độ soạn thảo văn bản."""
    return _skill("system_drafter.md")


def get_system_prompt_legal_qa() -> str:
    """System prompt cho chế độ hỏi luật."""
    return _skill("system_legal_qa.md")


# Aliases cho backward compatibility
SYSTEM_PROMPT          = get_system_prompt()
SYSTEM_PROMPT_LEGAL_QA = get_system_prompt_legal_qa()


# ============================================================
# 3. PRIORITY LABELS
# ============================================================

_LEGAL_PRIORITY: Dict[str, int] = {
    "LUẬT"      : 1,
    "PHÁP LỆNH" : 2,
    "NGHỊ ĐỊNH" : 3,
    "NGHỊ QUYẾT": 4,
}

_PRIORITY_BADGE: Dict[str, str] = {
    "LUẬT"      : "⚖️ [LUẬT — Hiệu lực cao nhất]",
    "PHÁP LỆNH" : "📜 [PHÁP LỆNH]",
    "NGHỊ ĐỊNH" : "📋 [NGHỊ ĐỊNH]",
    "NGHỊ QUYẾT": "📌 [NGHỊ QUYẾT]",
}

_DOC_ID_TYPE_RULES: List[Tuple[str, str]] = [
    ("NGHỊ ĐỊNH",  r'(^|[/\-_])ND($|[/\-_])|(^|[/\-_])NDCP($|[/\-_])'),
    ("NGHỊ QUYẾT", r'(^|[/\-_])NQ($|[/\-_])|(^|[/\-_])NQLT($|[/\-_])'),
    ("PHÁP LỆNH",  r'(^|[/\-_])PL($|[/\-_])'),
    ("LUẬT",       r'(^|[/\-_])QH\d*($|[/\-_])'),
]


def _fold_vietnamese(text: str) -> str:
    """Uppercase + bỏ dấu để nhận diện ổn định từ doc_id/số hiệu."""
    text = str(text or "").replace("Đ", "D").replace("đ", "d")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return text.upper()


def _normalize_legal_type(value: str) -> str:
    folded = _fold_vietnamese(value)
    for legal_type in _LEGAL_PRIORITY:
        if _fold_vietnamese(legal_type) == folded:
            return legal_type
    return ""


def _detect_type_from_doc_id(doc_id: str) -> str:
    folded = _fold_vietnamese(doc_id)
    for legal_type, pattern in _DOC_ID_TYPE_RULES:
        if re.search(pattern, folded):
            return legal_type
    return ""


def _get_legal_type(meta: Dict) -> str:
    """
    Lấy loại văn bản pháp luật chuẩn từ metadata.

    Ưu tiên:
    1. type_normalized nếu có và thuộc nhóm hỗ trợ priority badge.
    2. doc_id, vì thường chứa số hiệu dạng 170/2025/NĐ-CP__...
    3. source_doc_no/id nếu doc_id không có hoặc không suy ra được.
    """
    normalized = _normalize_legal_type(meta.get("type_normalized", ""))
    if normalized:
        return normalized

    for key in ("doc_id", "source_doc_no", "id"):
        detected = _detect_type_from_doc_id(meta.get(key, ""))
        if detected:
            return detected

    return "N/A"


def _priority_key(chunk: Dict) -> int:
    doc_type = _get_legal_type(chunk.get("metadata", {}))
    return _LEGAL_PRIORITY.get(doc_type, 99)


# ============================================================
# 4. FORM DETECTION
# ============================================================

# Từ khóa nhận diện Quyết định cá biệt (Form_02):
# Áp dụng trực tiếp lên cá nhân/tổ chức cụ thể — bổ nhiệm, kỷ luật, khen thưởng...
_QD_CA_BIET_PATTERN = re.compile(
    r"bo\s*nhiem|mien\s*nhiem|dieu\s*dong|tiep\s*nhan|khen\s*thuong"
    r"|ky\s*luat|nghi\s*viec|nghi\s*huu|cach\s*chuc|phe\s*duyet\s*nhan\s*su"
    r"|cho\s*phep|cap\s*phep|thu\s*hoi\s*giay\s*phep|nang\s*luong"
    r"|xep\s*luong|tuyen\s*dung|thoi\s*viec|giai\s*quyet",
    re.IGNORECASE,
)

# Từ khóa nhận diện Quyết định quy phạm/gián tiếp (Form_03):
# Ban hành văn bản quy phạm, quy chế, quy trình, tiêu chuẩn...
_QD_GIAN_TIEP_PATTERN = re.compile(
    r"ban\s*hanh\s*quy|ban\s*hanh\s*noi\s*quy|ban\s*hanh\s*tieu\s*chuan"
    r"|ban\s*hanh\s*dinh\s*muc|ban\s*hanh\s*phuong\s*an|ban\s*hanh\s*ke\s*hoach"
    r"|ban\s*hanh\s*quy\s*trinh|sua\s*doi\s*quy|bai\s*bo\s*quy"
    r"|ban\s*hanh\s*quy\s*che|ban\s*hanh\s*dinh\s*huong|ban\s*hanh\s*tieu\s*chi",
    re.IGNORECASE,
)

# Bảng keyword → Form ID (không bao gồm "quyet dinh" — xử lý riêng)
FORM_KEYWORD_MAP: List[Tuple[str, List[str]]] = [
    (r"nghi\s*quyet",                                               ["Form_01"]),
    # "quyet dinh" được xử lý động trong detect_form_ids()
    (r"to\s*trinh|trinh\s*duyet|trinh\s*phe\s*duyet",              ["Form_04"]),
    (r"bao\s*cao",                                                  ["Form_04"]),
    (r"chi\s*thi",                                                  ["Form_04"]),
    (r"thong\s*bao|thong\s*cao",                                    ["Form_04"]),
    (r"huong\s*dan",                                                ["Form_04"]),
    (r"ke\s*hoach|chuong\s*trinh|phuong\s*an|de\s*an|du\s*an",     ["Form_04"]),
    (r"quy\s*che|quy\s*dinh",                                       ["Form_04"]),
    (r"uy\s*quyen|giay\s*uy\s*quyen",                               ["Form_04"]),
    (r"phieu\s*gui|phieu\s*chuyen|phieu\s*bao",                     ["Form_04"]),
    (r"cong\s*van",                                                 ["Form_05"]),
    (r"cong\s*dien|khan\s*cap|thong\s*bao\s*khan",                  ["Form_06"]),
    (r"giay\s*moi|thu\s*moi|moi\s*hop|moi\s*hoi\s*nghi",           ["Form_07"]),
    (r"giay\s*gioi\s*thieu|gioi\s*thieu\s*can\s*bo",                ["Form_08"]),
    (r"bien\s*ban",                                                 ["Form_09"]),
    (r"giay\s*nghi\s*phep|nghi\s*phep|xin\s*nghi",                 ["Form_10"]),
]


def _detect_quyet_dinh_form(normalized_query: str) -> List[str]:
    """
    Phân biệt quyết định cá biệt (Form_02) vs quyết định quy phạm/gián tiếp (Form_03).

    Logic:
    - Có từ khóa nhân sự/cá nhân  → Form_02 (quyết định cá biệt)
    - Có từ khóa ban hành quy chế → Form_03 (quyết định gián tiếp)
    - Không rõ                     → [Form_02, Form_03] (fallback, để retriever chọn)

    Args:
        normalized_query: Query đã được normalize (lowercase, bỏ dấu)

    Returns:
        List Form ID phù hợp
    """
    if _QD_CA_BIET_PATTERN.search(normalized_query):
        return ["Form_02"]
    if _QD_GIAN_TIEP_PATTERN.search(normalized_query):
        return ["Form_03"]
    # Fallback: trả cả hai để retriever phân biệt theo vector similarity
    return ["Form_02", "Form_03"]


def detect_form_ids(query: str) -> List[str]:
    """
    Phân loại văn bản từ query người dùng, trả về danh sách Form ID phù hợp.

    Quyết định (quyet dinh) được xử lý riêng để phân biệt:
        - Form_02: Quyết định cá biệt (bổ nhiệm, kỷ luật, khen thưởng...)
        - Form_03: Quyết định gián tiếp/quy phạm (ban hành quy chế, quy trình...)
        - [Form_02, Form_03]: Fallback nếu không đủ tín hiệu

    Args:
        query: Yêu cầu soạn thảo của người dùng (có thể có dấu)

    Returns:
        List Form ID (vd: ["Form_02"] hoặc ["Form_02", "Form_03"])
        Trả về [] nếu không khớp loại văn bản nào.
    """
    from unidecode import unidecode  # pip install unidecode
    normalized = unidecode(query).lower()

    # Xử lý quyết định riêng vì có logic phân loại con
    if re.search(r"quyet\s*dinh", normalized, re.IGNORECASE):
        return _detect_quyet_dinh_form(normalized)

    # Các loại văn bản còn lại theo map tĩnh
    for pattern, form_ids in FORM_KEYWORD_MAP:
        if re.search(pattern, normalized, re.IGNORECASE):
            return form_ids

    return []


# ============================================================
# 5. CONTEXT FORMATTERS
# ============================================================

def _format_legal_context(legal_chunks: List[Dict]) -> str:
    if not legal_chunks:
        return "(Không tìm thấy điều khoản pháp luật liên quan.)"

    sorted_chunks = sorted(legal_chunks, key=_priority_key)
    parts: List[str] = []
    for idx, chunk in enumerate(sorted_chunks, start=1):
        meta     = chunk.get("metadata", {})
        doc_no   = meta.get("id", meta.get("source_doc_no", "N/A"))
        doc_type = _get_legal_type(meta)
        badge    = _PRIORITY_BADGE.get(doc_type, f"[{doc_type}]")
        doc_name = meta.get("doc_name", meta.get("name", "N/A"))
        article  = meta.get("article", "")
        text     = chunk.get("text", "").strip()
        score    = chunk.get("rerank_score", 0.0)
        eff_date = meta.get("effective_date", "")
        exp_date = meta.get("expiry_date", "")

        validity_note = ""
        if exp_date:
            validity_note = f"\n    Hết hiệu lực: {exp_date} ⚠️ — KIỂM TRA trước khi trích dẫn"
        elif eff_date:
            validity_note = f"\n    Có hiệu lực từ: {eff_date}"

        header = (
            f"{badge} [{idx}]\n"
            f"    Số hiệu : {doc_no}\n"
            f"    Tên VB  : {doc_name}\n"
            f"    Điều    : {article}"
            f"{validity_note}\n"
            f"    Độ liên quan: {score:.3f}"
        )
        parts.append(f"{header}\n\n{text}")

    return "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(parts) + "\n"


def _format_form_template(form_chunks: List[Dict]) -> str:
    if not form_chunks:
        return "(Không tìm thấy biểu mẫu phù hợp.)"

    form      = form_chunks[0]
    meta      = form.get("metadata", {})
    form_id   = meta.get("form_id", "N/A")
    form_type = meta.get("form_type", "N/A")
    purpose   = meta.get("purpose", "N/A")
    text      = form.get("text", "").strip()

    placeholders = sorted(set(re.findall(r"\{\{([^}]+)\}\}", text)))
    noi_dung_phs = [p for p in placeholders if "NOI_DUNG" in p]
    other_phs    = [p for p in placeholders if "NOI_DUNG" not in p]

    all_fields_list = "\n    ".join(f"- {ph}" for ph in other_phs) if other_phs \
                      else "(không xác định)"
    noi_dung_list   = "\n    ".join(f"- {ph}" for ph in noi_dung_phs) if noi_dung_phs \
                      else "(không có)"

    header = (
        f"Form ID   : {form_id}\n"
        f"Loại VB   : {form_type}\n"
        f"Mục đích  : {purpose}\n"
        f'\nField metadata (điền vào "fields"):\n'
        f"    {all_fields_list}\n"
        f'\nField thân văn bản — CŨNG điền vào "fields" (KHÔNG tạo key riêng):\n'
        f"    {noi_dung_list}"
    )
    return f"{header}\n\n{'─' * 60}\n\n{text}\n"


def _format_examples_context(example_chunks: List[Dict]) -> str:
    if not example_chunks:
        return "(Không tìm thấy ví dụ tương tự.)"

    parts: List[str] = []
    for idx, chunk in enumerate(example_chunks, start=1):
        meta     = chunk.get("metadata", {})
        scenario = meta.get("scenario", "N/A")
        text     = chunk.get("text", "").strip()
        parts.append(
            f"[Ví dụ {idx}]\n"
            f"Tình huống: {scenario}\n\n"
            f"{text}"
        )
    return "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(parts) + "\n"


def format_context(retrieved: Dict[str, List[Dict]]) -> Dict[str, str]:
    """
    Chuyển đổi kết quả retrieve_all() thành các khối văn bản cho prompt.

    Args:
        retrieved: Dict từ retrieve_all() với keys "legal", "form", "examples"

    Returns:
        Dict với keys: "legal_context", "form_template", "few_shot_examples"
    """
    return {
        "legal_context"     : _format_legal_context(retrieved.get("legal", [])),
        "form_template"     : _format_form_template(retrieved.get("form", [])),
        "few_shot_examples" : _format_examples_context(retrieved.get("examples", [])),
    }


# ============================================================
# 6. NOI_DUNG STRUCTURE HINT (từ skill file)
# ============================================================

_FORM_TYPE_HINT_KEY: Dict[str, str] = {
    "công văn"         : "CÔNG VĂN",
    "quyết định"       : "QUYẾT ĐỊNH",
    "quyết định cá biệt"  : "QUYẾT ĐỊNH",
    "quyết định gián tiếp": "QUYẾT ĐỊNH",
    "tờ trình"         : "TỜ TRÌNH",
    "biên bản"         : "BIÊN BẢN",
    "giấy nghỉ phép"   : "GIẤY NGHỈ PHÉP",
}

_SECTION_HEADER_MAP: Dict[str, str] = {
    "CÔNG VĂN"      : "## CÔNG VĂN",
    "QUYẾT ĐỊNH"    : "## QUYẾT ĐỊNH",
    "TỜ TRÌNH"      : "## TỜ TRÌNH",
    "BIÊN BẢN"      : "## BIÊN BẢN",
    "GIẤY NGHỈ PHÉP": "## GIẤY NGHỈ PHÉP",
    "MẶC ĐỊNH"      : "## MẶC ĐỊNH",
}


def _get_noi_dung_hint(form_type: str) -> str:
    """
    Trích section hướng dẫn cấu trúc NOI_DUNG_* tương ứng từ skill file.
    Fallback về section MẶC ĐỊNH nếu không khớp loại văn bản.
    """
    skill_text = _skill("noi_dung_structure.md")

    normalized  = form_type.strip().lower()
    section_key = _FORM_TYPE_HINT_KEY.get(normalized, "MẶC ĐỊNH")
    header      = _SECTION_HEADER_MAP[section_key]

    start = skill_text.find(header)
    if start == -1:
        default_header = _SECTION_HEADER_MAP["MẶC ĐỊNH"]
        start = skill_text.find(default_header)
        if start == -1:
            return skill_text

    next_section = skill_text.find("\n## ", start + len(header))
    section_text = skill_text[start:next_section].strip() if next_section != -1 \
                   else skill_text[start:].strip()

    lines = section_text.split("\n")
    content_lines = [l for l in lines[1:] if l.strip() or lines.index(l) > 1]
    return "\n".join(content_lines).strip()


def _extract_form_type_from_template(form_tmpl: str) -> str:
    m = re.search(r"Loại VB\s*:\s*(.+)", form_tmpl)
    return m.group(1).strip() if m else ""


# ============================================================
# 7. USER PROMPT BUILDERS
# ============================================================

def build_user_prompt(
    query: str,
    context: Dict[str, str],
    extra_instructions: Optional[str] = None,
) -> str:
    """
    Xây dựng user prompt yêu cầu LLM trả về JSON thuần túy với DUY NHẤT key "fields".

    Args:
        query              : Yêu cầu soạn thảo của người dùng
        context            : Dict từ format_context()
        extra_instructions : Hướng dẫn bổ sung (ngày ký, người ký, số VB...)

    Returns:
        Chuỗi prompt hoàn chỉnh
    """
    legal_ctx = context["legal_context"]
    form_tmpl = context["form_template"]
    few_shot  = context["few_shot_examples"]

    extra_block = ""
    if extra_instructions:
        extra_block = (
            "\n## YÊU CẦU BỔ SUNG\n"
            + extra_instructions.strip()
            + "\n"
        )

    all_placeholders = sorted(set(re.findall(r"\{\{([^}]+)\}\}", form_tmpl)))
    noi_dung_fields  = [p for p in all_placeholders if "NOI_DUNG" in p]
    other_fields     = [p for p in all_placeholders if "NOI_DUNG" not in p]

    all_field_lines = [f'    "{p}": "..."' for p in other_fields]
    for p in noi_dung_fields:
        hint_short = "(nội dung thân văn bản — xem hướng dẫn cấu trúc quy tắc 2)"
        all_field_lines.append(f'    "{p}": "{hint_short}"')
    fields_example = ",\n".join(all_field_lines) if all_field_lines \
                     else '    "(không có field)"'

    form_type     = _extract_form_type_from_template(form_tmpl)
    noi_dung_hint = _get_noi_dung_hint(form_type)

    if noi_dung_fields:
        noi_dung_rule = (
            '2. CÁC FIELD THÂN VĂN BẢN — điền trực tiếp vào key "fields":\n'
            + "".join(f"   • {p}\n" for p in noi_dung_fields)
            + "   Viết nội dung theo cấu trúc sau:\n"
            + "   " + noi_dung_hint.replace("\n", "\n   ")
            + '\n   KHÔNG tạo key riêng ngoài "fields" cho các field này.'
        )
    else:
        noi_dung_rule = (
            '2. Form này không có field NOI_DUNG_* — điền toàn bộ thông tin vào "fields".'
        )

    legal_block_header = _skill("block_legal_context.md")
    json_output_block  = _skill("block_json_output.md")

    json_output_final = re.sub(
        r'```json\n\{[\s\S]*?\}\n```',
        f'{{\n  "fields": {{\n{fields_example}\n  }}\n}}',
        json_output_block,
    )
    json_output_final = re.sub(
        r'(### Quy tắc 2.*?)(?=### Quy tắc 3)',
        noi_dung_rule + "\n\n",
        json_output_final,
        flags=re.DOTALL,
    )

    prompt = (
        "## YÊU CẦU SOẠN THẢO\n"
        + query.strip()
        + "\n"
        + extra_block
        + "\n"
        + "=" * 64 + "\n"
        + legal_block_header + "\n"
        + legal_ctx
        + "\n"
        + "=" * 64 + "\n"
        + "## FORM_TEMPLATE — Biểu mẫu và danh sách FIELD_NAME\n"
        + form_tmpl
        + "\n"
        + "=" * 64 + "\n"
        + "## FEW_SHOT_EXAMPLES — Ví dụ tham khảo văn phong và cách điền\n"
        + few_shot
        + "\n"
        + "=" * 64 + "\n"
        + json_output_final
    )

    return prompt


def build_legal_qa_user_prompt(
    query: str,
    legal_context: str,
    extra_instructions: Optional[str] = None,
) -> str:
    """
    Xây dựng user prompt cho chế độ hỏi luật (không cần form/example).

    Args:
        query             : Câu hỏi pháp luật của người dùng
        legal_context     : Chuỗi đã format từ _format_legal_context()
        extra_instructions: Hướng dẫn bổ sung tuỳ chọn

    Returns:
        Chuỗi prompt hoàn chỉnh
    """
    extra_block = ""
    if extra_instructions:
        extra_block = (
            "\n## YÊU CẦU BỔ SUNG\n"
            + extra_instructions.strip()
            + "\n"
        )

    legal_block_header = _skill("block_legal_context.md")
    answer_block       = _skill("block_legal_qa_answer.md")

    prompt = (
        "## CÂU HỎI PHÁP LUẬT\n"
        + query.strip()
        + "\n"
        + extra_block
        + "\n"
        + "=" * 64 + "\n"
        + legal_block_header + "\n"
        + legal_context
        + "\n"
        + "=" * 64 + "\n"
        + answer_block
    )
    return prompt


# ============================================================
# 8. MESSAGES BUILDERS
# ============================================================

def build_messages(
    query: str,
    retrieved: Dict[str, List[Dict]],
    extra_instructions: Optional[str] = None,
) -> List[Dict[str, str]]:
    """
    Xây dựng danh sách messages cho API (OpenAI / Anthropic chat format).

    Args:
        query              : Yêu cầu soạn thảo (tiếng Việt)
        retrieved          : Dict từ retrieve_all()
        extra_instructions : Hướng dẫn bổ sung (ngày ký, người ký, số VB...)

    Returns:
        [{"role": "system", "content": ...}, {"role": "user", "content": ...}]
    """
    ctx         = format_context(retrieved)
    user_prompt = build_user_prompt(query, ctx, extra_instructions)
    return [
        {"role": "system", "content": get_system_prompt()},
        {"role": "user",   "content": user_prompt},
    ]


def build_legal_qa_messages(
    query: str,
    retrieved_legal: List[Dict],
    extra_instructions: Optional[str] = None,
) -> List[Dict[str, str]]:
    """
    Xây dựng messages cho chế độ HỎI LUẬT (không cần form/example).

    Args:
        query            : Câu hỏi pháp luật (tiếng Việt)
        retrieved_legal  : List[Dict] từ retrieve_legal() với top_k=5
        extra_instructions: Hướng dẫn bổ sung tuỳ chọn

    Returns:
        [{"role": "system", "content": ...}, {"role": "user", "content": ...}]
    """
    legal_ctx   = _format_legal_context(retrieved_legal)
    user_prompt = build_legal_qa_user_prompt(query, legal_ctx, extra_instructions)
    return [
        {"role": "system", "content": get_system_prompt_legal_qa()},
        {"role": "user",   "content": user_prompt},
    ]


# ============================================================
# 9. JSON PARSER & VALIDATOR
# ============================================================

def parse_llm_json(raw: str) -> Dict:
    """
    Parse JSON từ LLM response. Xử lý các trường hợp LLM bọc code fence
    hoặc thêm text thừa trước/sau JSON.

    Thứ tự thử:
    1. Extract nội dung trong ```json ... ``` hoặc ``` ... ```
    2. Parse trực tiếp toàn bộ text (sau khi strip)
    3. Tìm JSON object lớn nhất trong text (từ { đầu đến } cuối)

    Args:
        raw: Toàn bộ string trả về từ LLM

    Returns:
        Dict đã chuẩn hóa với key duy nhất "fields" (dict chứa tất cả placeholder)

    Raises:
        ValueError nếu không tìm được JSON hợp lệ
    """
    text = raw.strip()

    # ── Bước 1: Extract content bên trong code fence (```json ... ``` hoặc ``` ... ```) ──
    # Dùng re.DOTALL để . khớp cả newline; non-greedy để lấy fence đầu tiên
    fence_match = re.search(r"```(?:json)?\s*\n?([\s\S]+?)\n?\s*```", text)
    if fence_match:
        candidate = fence_match.group(1).strip()
        try:
            return _validate_structure(json.loads(candidate))
        except json.JSONDecodeError:
            pass  # tiếp tục fallback

    # ── Bước 2: Parse trực tiếp (LLM trả JSON thuần, không có fence) ──
    try:
        return _validate_structure(json.loads(text))
    except json.JSONDecodeError:
        pass

    # ── Bước 3: Tìm JSON object lớn nhất — từ { đầu tiên đến } cuối cùng ──
    # Dùng rfind('}') thay vì greedy regex để tránh cắt sót khi value có ngoặc lồng
    start = text.find("{")
    end   = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return _validate_structure(json.loads(text[start : end + 1]))
        except json.JSONDecodeError:
            pass

    raise ValueError(
        "Không parse được JSON từ LLM output.\n"
        f"200 ký tự đầu: {raw[:200]!r}"
    )


def _validate_structure(data: Dict) -> Dict:
    """
    Chuẩn hóa và validate cấu trúc JSON từ LLM.
    Đảm bảo chỉ có key "fields" chứa toàn bộ placeholder (kể cả NOI_DUNG_*).

    Nếu LLM vẫn trả về key "noi_dung" riêng (trái với hướng dẫn), tự động
    gộp nó vào field NOI_DUNG tương ứng trong "fields" hoặc tạo key dự phòng.
    """
    if not isinstance(data, dict):
        raise ValueError(f"JSON root phải là object, nhận: {type(data).__name__}")

    # Recover nếu LLM trả flat dict (bỏ lớp "fields")
    if "fields" not in data:
        noi_dung_val = data.pop("noi_dung", "")
        data = {"fields": data, "_orphan_noi_dung": noi_dung_val}

    if not isinstance(data["fields"], dict):
        data["fields"] = {}

    # Gộp key "noi_dung" lạc vào "fields" nếu LLM vẫn tạo
    orphan = data.pop("_orphan_noi_dung", None) or data.pop("noi_dung", None)
    if orphan:
        noi_dung_keys = sorted(k for k in data["fields"] if "NOI_DUNG" in k.upper())
        if noi_dung_keys:
            first_key = noi_dung_keys[0]
            if not data["fields"].get(first_key):
                data["fields"][first_key] = str(orphan)
            else:
                data["fields"][first_key] = (
                    data["fields"][first_key] + "\n\n" + str(orphan)
                ).strip()
        else:
            data["fields"]["NOI_DUNG"] = str(orphan)

    # Xóa các key ngoài "fields" nếu LLM vô tình thêm
    for k in [k for k in data if k != "fields"]:
        data.pop(k, None)

    # Ép value về string, loại None
    data["fields"] = {
        k: str(v) if v is not None else ""
        for k, v in data["fields"].items()
    }

    return data


# ============================================================
# 10. WORD TEMPLATE FILLER
# ============================================================

def fill_word_template(
    template_path: str,
    parsed: Dict,
    output_path: str,
) -> str:
    """
    Điền giá trị từ JSON đã parse vào Word template (.docx).

    Placeholder trong .docx phải có dạng {{FIELD_NAME}}.
    Tất cả field (kể cả NOI_DUNG_*) lấy từ parsed["fields"].

    Args:
        template_path: Đường dẫn .docx template (có chứa {{FIELD_NAME}})
        parsed       : Dict từ parse_llm_json()
        output_path  : Đường dẫn file .docx kết quả

    Returns:
        output_path (string)

    Raises:
        ImportError nếu python-docx chưa được cài
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx chưa được cài. Chạy: pip install python-docx"
        )

    mapping: Dict[str, str] = {
        str(k).strip(): str(v) if v is not None else ""
        for k, v in dict(parsed.get("fields", {})).items()
    }

    def _normalize_field_key(name: str) -> str:
        return re.sub(r"[^A-Za-z0-9]+", "", name).upper()

    normalized_mapping: Dict[str, str] = {}
    for key in mapping:
        normalized_mapping.setdefault(_normalize_field_key(key), key)

    placeholder_pattern = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")
    unresolved_placeholders: set[str] = set()

    def _resolve_placeholder(raw_key: str) -> Optional[str]:
        key = raw_key.strip()
        if key in mapping:
            return mapping[key]
        alias = normalized_mapping.get(_normalize_field_key(key))
        if alias is not None:
            return mapping[alias]
        unresolved_placeholders.add(key)
        return None

    def _replace_text(text: str) -> tuple[str, bool]:
        changed = False

        def _sub(match: re.Match) -> str:
            nonlocal changed
            value = _resolve_placeholder(match.group(1))
            if value is None:
                return match.group(0)
            changed = True
            return value

        return placeholder_pattern.sub(_sub, text), changed

    doc           = Document(template_path)
    replace_count = 0

    def _replace_para(para) -> bool:
        full_text = "".join(r.text for r in para.runs)
        if not placeholder_pattern.search(full_text):
            return False

        changed = False
        for run in para.runs:
            new_text, run_changed = _replace_text(run.text)
            if run_changed:
                run.text = new_text
                changed = True

        # Fallback: placeholder bị Word split thành nhiều run
        if not changed and para.runs:
            new_full, full_changed = _replace_text(full_text)
            if full_changed:
                first_run            = para.runs[0]
                first_run.text       = new_full
                first_run.bold       = first_run.bold
                first_run.italic     = first_run.italic
                first_run.font.name  = first_run.font.name
                first_run.font.size  = first_run.font.size
                for r in para.runs[1:]:
                    r.text = ""
                changed = True

        return changed

    for para in doc.paragraphs:
        if _replace_para(para):
            replace_count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_para(para):
                        replace_count += 1

    for section in doc.sections:
        for para in section.header.paragraphs:
            if _replace_para(para):
                replace_count += 1
        for para in section.footer.paragraphs:
            if _replace_para(para):
                replace_count += 1

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    if unresolved_placeholders:
        unresolved = sorted(unresolved_placeholders)
        preview = ", ".join(unresolved[:20])
        more = "" if len(unresolved) <= 20 else f", ... (+{len(unresolved) - 20})"
        print(f"[WARN] Placeholder chưa có dữ liệu: {preview}{more}")

    print(f"[OK] Đã lưu Word: {output_path}  ({replace_count} đoạn được thay thế)")
    return output_path


# ============================================================
# 11. UTILITIES — tương thích với generatePrompt.py
# ============================================================

def find_unfilled_placeholders(text: str) -> List[str]:
    """Tìm {{FIELD_NAME}} còn sót lại trong văn bản."""
    return re.findall(r"\{\{([^}]+)\}\}", text)


def get_context_stats(retrieved: Dict[str, List[Dict]]) -> Dict[str, int]:
    """Thống kê số chunks retrieved theo loại."""
    return {
        "legal"   : len(retrieved.get("legal", [])),
        "form"    : len(retrieved.get("form", [])),
        "examples": len(retrieved.get("examples", [])),
    }


def is_legal_qa_mode(retrieved: Dict[str, List[Dict]]) -> bool:
    """Kiểm tra dict retrieved có phải từ chế độ hỏi luật không."""
    return bool(retrieved.get("legal")) and not retrieved.get("form") and not retrieved.get("examples")


def reload_skills() -> None:
    _skill_cache.clear()
    print(f"[OK] Đã xóa skill cache — {len(list(SKILLS_DIR.glob('*.md')))} file sẽ được reload khi dùng.")
